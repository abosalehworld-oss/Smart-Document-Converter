"""
Word Generator - إنشاء ملفات Word
===================================
إنشاء مستندات Word (.docx) مع دعم RTL أصلي.
⚠️ لا نستخدم arabic-reshaper هنا - Word يدعم RTL أصلاً.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import logging
import os

from app.utils.arabic_utils import detect_language, normalize_arabic_text

logger = logging.getLogger(__name__)


class WordGenerator:
    """
    مولد مستندات Word مع دعم كامل للعربي (RTL).
    """
    
    def __init__(
        self,
        font_name: str = 'Simplified Arabic',
        font_size: int = 14,
        english_font: str = 'Arial'
    ):
        self.font_name = font_name
        self.font_size = font_size
        self.english_font = english_font
        self._doc = None
        self._page_count = 0
    
    def create_document(self):
        """إنشاء مستند Word جديد."""
        self._doc = Document()
        self._page_count = 0
        
        # ضبط إعدادات المستند الافتراضية
        self._setup_default_styles()
        
        # ضبط هوامش الصفحة
        for section in self._doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def _setup_default_styles(self):
        """ضبط الأنماط الافتراضية للمستند."""
        style = self._doc.styles['Normal']
        
        # خط افتراضي
        font = style.font
        font.name = self.english_font
        font.size = Pt(self.font_size)
        
        # خط عربي (Complex Script font)
        rpr = style.element.get_or_add_rPr()
        cs_font = rpr.find(qn('w:rFonts'))
        if cs_font is None:
            cs_font = parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="{self.font_name}"/>')
            rpr.append(cs_font)
        else:
            cs_font.set(qn('w:cs'), self.font_name)
        
        # حجم خط Complex Script
        cs_size = rpr.find(qn('w:szCs'))
        if cs_size is None:
            cs_size = parse_xml(f'<w:szCs {nsdecls("w")} w:val="{self.font_size * 2}"/>')
            rpr.append(cs_size)
        else:
            cs_size.set(qn('w:val'), str(self.font_size * 2))
    
    def add_page_text(self, text: str, page_number: int = None):
        """
        إضافة نص صفحة كاملة للمستند.
        يكتشف اتجاه النص تلقائياً (عربي RTL / إنجليزي LTR).
        
        Args:
            text: النص المراد إضافته
            page_number: رقم الصفحة (اختياري، للعرض)
        """
        if self._doc is None:
            self.create_document()
        
        # إضافة فاصل صفحة (إلا للصفحة الأولى)
        if self._page_count > 0:
            self._doc.add_page_break()
        
        # تطبيع النص
        text = normalize_arabic_text(text)
        
        if not text.strip():
            p = self._doc.add_paragraph("[صفحة فارغة / Empty page]")
            self._page_count += 1
            return
        
        # تقسيم لفقرات
        paragraphs = text.split('\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            
            if not para_text:
                # سطر فارغ
                self._doc.add_paragraph('')
                continue
            
            # إنشاء فقرة
            paragraph = self._doc.add_paragraph()
            
            # كشف اللغة وضبط الاتجاه
            lang = detect_language(para_text)
            
            if lang in ('ar', 'mixed'):
                # نص عربي أو مختلط - RTL
                self._set_paragraph_rtl(paragraph)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                # نص إنجليزي - LTR
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # إضافة النص
            run = paragraph.add_run(para_text)
            
            # ضبط الخط
            if lang in ('ar', 'mixed'):
                run.font.rtl = True
                # خط Complex Script للعربي
                rpr = run._element.get_or_add_rPr()
                cs_font = parse_xml(
                    f'<w:rFonts {nsdecls("w")} w:cs="{self.font_name}"/>'
                )
                rpr.append(cs_font)
                cs_size = parse_xml(
                    f'<w:szCs {nsdecls("w")} w:val="{self.font_size * 2}"/>'
                )
                rpr.append(cs_size)
            else:
                run.font.name = self.english_font
            
            run.font.size = Pt(self.font_size)
        
        self._page_count += 1
    
    def _set_paragraph_rtl(self, paragraph):
        """ضبط اتجاه الفقرة RTL."""
        pPr = paragraph._element.get_or_add_pPr()
        
        # BiDi property
        bidi = pPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = parse_xml(f'<w:bidi {nsdecls("w")}/>')
            pPr.append(bidi)
    
    def add_header(self, text: str, level: int = 1):
        """إضافة عنوان."""
        if self._doc is None:
            self.create_document()
        
        heading = self._doc.add_heading(text, level=level)
        
        if detect_language(text) in ('ar', 'mixed'):
            self._set_paragraph_rtl(heading)
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def save(self, output_path: str) -> str:
        """
        حفظ المستند.
        
        Args:
            output_path: مسار الحفظ
        
        Returns:
            المسار الفعلي للملف المحفوظ
        """
        if self._doc is None:
            raise ValueError("لم يتم إنشاء مستند")
        
        # التأكد من امتداد .docx
        if not output_path.lower().endswith('.docx'):
            output_path += '.docx'
        
        # التأكد من وجود المجلد
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        try:
            self._doc.save(output_path)
            logger.info(f"تم حفظ المستند: {self._page_count} صفحة")
            return output_path
        except Exception as e:
            logger.error(f"فشل حفظ المستند: {type(e).__name__}")
            raise
    
    def reset(self):
        """إعادة تعيين المولد لمستند جديد."""
        self._doc = None
        self._page_count = 0
    
    @property
    def page_count(self) -> int:
        """عدد الصفحات المضافة."""
        return self._page_count
